import re
import os
from typing import Optional

# PDF parsing
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    from pdfminer.layout import LAParams
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ??? Section header patterns ????????????????????????????????????????????????

SECTION_PATTERNS = {
    'summary': re.compile(
        r'^\s*(professional\s+)?summary|objective|profile|about\s+me|career\s+objective\s*$',
        re.IGNORECASE | re.MULTILINE
    ),
    'education': re.compile(
        r'^\s*education|academic|qualifications|degrees?\s*$',
        re.IGNORECASE | re.MULTILINE
    ),
    'experience': re.compile(
        r'^\s*(work\s+)?(experience|history)|employment|career\s+history|professional\s+experience\s*$',
        re.IGNORECASE | re.MULTILINE
    ),
    'skills': re.compile(
        r'^\s*(technical\s+)?skills|competencies|technologies|expertise|proficiencies\s*$',
        re.IGNORECASE | re.MULTILINE
    ),
    'projects': re.compile(
        r'^\s*projects?|personal\s+projects?|key\s+projects?\s*$',
        re.IGNORECASE | re.MULTILINE
    ),
    'certifications': re.compile(
        r'^\s*certifications?|certificates?|licenses?|accreditations?\s*$',
        re.IGNORECASE | re.MULTILINE
    ),
    'achievements': re.compile(
        r'^\s*achievements?|accomplishments?|awards?|honors?\s*$',
        re.IGNORECASE | re.MULTILINE
    ),
    'languages': re.compile(
        r'^\s*languages?\s*$',
        re.IGNORECASE | re.MULTILINE
    ),
}


def extract_text_from_pdf(filepath: str) -> str:
    """Extract plain text from a PDF file using pdfminer.six."""
    if not PDF_AVAILABLE:
        raise RuntimeError('pdfminer.six is not installed')

    try:
        laparams = LAParams(
            line_overlap=0.5,
            char_margin=2.0,
            line_margin=0.5,
            word_margin=0.1,
            boxes_flow=0.5,
            detect_vertical=False,
        )
        text = pdf_extract_text(filepath, laparams=laparams)
        return text or ''
    except Exception as e:
        raise RuntimeError(f'PDF extraction failed: {str(e)}') from e


def extract_text_from_docx(filepath: str) -> str:
    """Extract plain text from a DOCX file using python-docx."""
    if not DOCX_AVAILABLE:
        raise RuntimeError('python-docx is not installed')

    try:
        doc = Document(filepath)
        paragraphs = []

        # Extract from paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)

        return '\n'.join(paragraphs)
    except Exception as e:
        raise RuntimeError(f'DOCX extraction failed: {str(e)}') from e


def _detect_sections(text: str) -> dict:
    """
    Detect resume sections by scanning for header keywords.
    Returns a dict mapping section_name -> section_text.
    """
    lines = text.split('\n')
    sections = {}
    current_section: Optional[str] = None
    current_lines = []

    for line in lines:
        stripped = line.strip()
        matched_section = None

        # Only consider short lines (likely headers, not paragraph text)
        if len(stripped) <= 60:
            for section_name, pattern in SECTION_PATTERNS.items():
                if pattern.match(stripped):
                    matched_section = section_name
                    break

        if matched_section:
            # Save previous section
            if current_section and current_lines:
                sections[current_section] = '\n'.join(current_lines).strip()
            current_section = matched_section
            current_lines = []
        else:
            if current_section:
                current_lines.append(line)

    # Save last section
    if current_section and current_lines:
        sections[current_section] = '\n'.join(current_lines).strip()

    return sections


def parse_resume(filepath: str) -> dict:
    """
    Parse a resume file (PDF or DOCX).
    Returns:
        {
            'raw_text': str,
            'sections': {
                'summary': str,
                'education': str,
                'experience': str,
                'skills': str,
                'projects': str,
                'certifications': str,
                ...
            }
        }
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f'File not found: {filepath}')

    ext = filepath.rsplit('.', 1)[-1].lower()

    if ext == 'pdf':
        raw_text = extract_text_from_pdf(filepath)
    elif ext == 'docx':
        raw_text = extract_text_from_docx(filepath)
    else:
        raise ValueError(f'Unsupported file type: {ext}')

    # Normalize whitespace but preserve line breaks
    raw_text = re.sub(r'\r\n', '\n', raw_text)
    raw_text = re.sub(r'\r', '\n', raw_text)
    raw_text = re.sub(r'[ \t]+', ' ', raw_text)  # collapse horizontal whitespace
    raw_text = re.sub(r'\n{3,}', '\n\n', raw_text)  # collapse excess blank lines
    raw_text = raw_text.strip()

    # Detect sections
    sections = _detect_sections(raw_text)

    return {
        'raw_text': raw_text,
        'sections': sections,
    }
